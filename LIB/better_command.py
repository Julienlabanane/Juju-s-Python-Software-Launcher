import LIB.system_sniffer as system_sniffer
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from pathlib import Path

search_timenow = datetime.now()
timenow = search_timenow.strftime("%d-%m-%Y-%H_%M_%S.%f")
log_dir = Path("LOG")
log_dir.mkdir(parents=True, exist_ok=True)
log_file_name = log_dir / f"JPSL_log_{timenow}.docx"
log_file = Document()
log_file.add_heading("JPSL Log File", 0)
log_file.add_paragraph("Log created on: " + timenow)
system_sniffer_result = (
    ('CPU', system_sniffer.get_cpu_name()),
    ('RAM', system_sniffer.get_ram()),
    ('OS', system_sniffer.get_os_info()),
)
table = log_file.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Model'
for id, desc in system_sniffer_result:
    row_cells = table.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = desc

def better_command(txt, context):
    yellow_warning = '\033[33m'
    error_warning = '\033[91m'
    succes = '\033[32m'
    base = '\033[0m'
    libname = "Juju's Python Better Command"
    libver = '1'
    entry_time = datetime.now().strftime("%d-%m-%Y %H:%M:%S.%f")
    if context == "hello":
        print(yellow_warning + "◼" + error_warning + "◼" + succes + "◼" + base + "  " + libname + " v" + libver)
    if context == "error":
        print(error_warning + "[" + entry_time + "] → " + txt)
        error_indoc = log_file.add_paragraph(entry_time + " → " + txt + "\n")
        error_indoc.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    if context == "warning":
        print(yellow_warning + "[" + entry_time + "] → " + txt)
        warning_indoc = log_file.add_paragraph(entry_time + " → " + txt + "\n")
        warning_indoc.runs[0].font.color.rgb = RGBColor(255, 255, 0)
    if context == "success":
        print(succes + "[" + entry_time + "] → " + txt)
        success_indoc = log_file.add_paragraph(entry_time + " → " + txt + "\n")
        success_indoc.runs[0].font.color.rgb = RGBColor(0, 255, 0)
    if context == "info":
        print(base + "[" + entry_time + "] → " + txt)
        log_file.add_paragraph("[INFO] " + entry_time + " → " + txt + "\n")
    log_file.save(str(log_file_name))

better_command("", "hello")